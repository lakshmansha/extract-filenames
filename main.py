import os
from docx import Document

def extract_file_names_to_word(folder_path, output_word_path, columns=3):
    """
    Extracts file names from a folder and writes them into a Word document in tabular format with multiple columns.

    :param folder_path: Path to the folder containing files.
    :param output_word_path: Path to save the Word document.
    :param columns: Number of columns in the table.
    """
    try:
        # Check if the folder exists
        if not os.path.exists(folder_path):
            print(f"The folder '{folder_path}' does not exist.")
            return

        # List all files in the folder
        file_names = [file for file in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, file))]

        # Remove prefix 'MOH_' and suffix '.jpg' from file names
        processed_names = [file[4:-4] for file in file_names]

        # Create a new Word document
        document = Document()
        document.add_heading('Files on the Folder', level=1)

        # Add a table to the document
        rows = (len(processed_names) + columns - 1) // columns  # Calculate required rows
        table = document.add_table(rows=rows + 1, cols=columns)
        table.style = 'Table Grid'

        # Add header row
        # for i in range(columns):
        #     header_cell = table.rows[0].cells[i]
        #     header_cell.text = f'File Name {i + 1}'
        
        # Remove table borders
        # for row in table.rows:
        #     for cell in row.cells:
        #         cell._element.get_or_add_tcPr().append(Document().element.xpath('.//w:tcBorders')[0].clear())


        # Add file names to the table
        for index, file_name in enumerate(processed_names):
            row_idx = index // columns
            col_idx = index % columns
            table.rows[row_idx].cells[col_idx].text = f"{index + 1}.)  {file_name}"

        # Save the document
        document.save(output_word_path)
        print(f"Word document saved successfully at '{output_word_path}'.")

    except Exception as e:
        print(f"An error occurred: {e}")

# Example usage
folder_path = "input"  # Replace with the path to your folder
output_word_path = "output_file.docx"  # Replace with the desired Word document name
extract_file_names_to_word(folder_path, output_word_path, columns=4)
